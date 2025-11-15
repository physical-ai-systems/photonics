from docx import Document
from docx.shared import Inches
from Utils.Utils import get_images_from_folder, get_data_and_config_paths, load_yaml
import os

class ExportUtils:

    @staticmethod
    def export_images_to_word(experiment_path: str, physical_quantity: dict, output_path: str, keep_old: bool = True, open_file: bool = True):
        image_paths = ExportUtils.get_images_related_to_physical_quantity(experiment_path, physical_quantity)
        # Create a new Word document
        if os.path.exists(output_path) and keep_old:
            doc = Document(output_path)
        else:
            doc = Document()

        # Loop through the image paths and insert them into the document
        for image_path in image_paths:
            # Add a new paragraph
            doc.add_paragraph()

            # Add the image to the paragraph
            doc.paragraphs[-1].add_run().add_picture(image_path, width=Inches(4))

        # Save the document
        doc.save(output_path)


        # Open the document
        if open_file:
        
            try :
                # open the document in microsoft word in mac
                if os.name == 'posix':
                    os.system(f'open {output_path}')
                # open the document in microsoft word in windows
                elif os.name == 'nt':
                    os.system(f'start winword {output_path}')
                else:
                    os.system(f'start {output_path}')
            except:
                print("Error in opening the document")
                

    @staticmethod
    def get_images_related_to_physical_quantity(experiment_path: str, physical_quantity: dict):
        image_paths = []

        for key in physical_quantity:
            _ , config_path = get_data_and_config_paths(experiment_path, key)
            config, images_dir = load_yaml(config_path), 'images'
            all_images = get_images_from_folder(os.path.join(experiment_path, images_dir))

            ExportUtils.get_images_related_to_word_insertion(config, all_images, image_paths, key)


        return image_paths

    @staticmethod
    def get_images_related_to_word_insertion(config, all_images, image_paths, main_key):
        word_sub_key = 'add_to_word'
        image_indices = []
        image_keys    = []
        images_keys_based_on_paths  = [os.path.basename(x).split(".")[-2] for x in all_images]
        for key in config:
            if word_sub_key in  config[key]:
                if config[key][word_sub_key] is not None:
                    image_indices.append(config[key][word_sub_key])
                    image_keys.append(key)
        # sort the images based on the indices
        image_keys = [x for _, x in sorted(zip(image_indices, image_keys))]
        # Adding main key to the image keys
        image_keys = [ main_key + '_' + image_key for image_key in image_keys]

        #Check if the image exists and add it to the all images
        for key in image_keys:
            if key in images_keys_based_on_paths:
                image_paths.append(all_images[images_keys_based_on_paths.index(key)])
            else:
                print(f"Image {key} not found in the folder")
                    
                
                


